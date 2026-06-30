#!/usr/bin/env python3
"""Add comparison section to optical full chain doc."""
from docx import Document
from docx.shared import Pt, RGBColor

doc = Document(r'C:\Users\Duke Wang\.openclaw\workspace\光通信全产业链_细分赛道与标的全景_20260604.docx')

# Add comparison section
doc.add_heading('九、估值对比：光通信 vs 存储芯片', level=1)

doc.add_paragraph('为什么光通信PE比存储高5倍？五个结构性因素：')

factors = [
    ('1. 成长阶段不同：', '光通信(800G→1.6T)处于爆发前期，存储已到周期中期。光通信需求曲线更陡峭，市场在定价未来2-3年的加速增长。'),
    ('2. 资产模型不同：', '光模块轻资产(产线几亿起步)，存储重资产(晶圆厂100-200亿)。轻资产企业天然享有更高PE估值。'),
    ('3. 全球地位不同：', '中国光模块全球份额60%+，中际旭创全球第一。存储芯片仍在追赶(长存NAND 14%、长鑫DRAM 8-12%)。全球龙头vs国产追赶者的PE溢价天然存在。'),
    ('4. 景气确定性不同：', '光通信直接挂钩英伟达/谷歌/Meta资本开支(持续上修)。存储受宏观+消费电子双重影响，周期性更强、不确定性更大。'),
    ('5. 市场叙事不同：', '光通信="AI血管"新故事，想象空间大。存储="老周期涨价"，历史包袱重。投资者结构和情绪差异显著。'),
]

for title, desc in factors:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    p.add_run(desc)

p = doc.add_paragraph()
run = p.add_run('结论：')
run.bold = True
p.add_run('光通信高PE有一定合理性(轻资产+全球龙头+成长早期)，但当前溢价已过度。非龙头PE普遍200-2000x属严重泡沫。中际旭创Forward PE 58x(基于FY26E净利90亿)是光通信中最有安全边际的标的，与存储中德明利(PE 38x)的角色类似。')

doc.save(r'C:\Users\Duke Wang\.openclaw\workspace\光通信全产业链_细分赛道与标的全景_20260604.docx')
print('Updated with comparison section')
