#!/usr/bin/env python3
"""Generate comprehensive optical communication industry chain Word report."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)

title = doc.add_heading('', level=0)
run = title.add_run('光通信全产业链 — 细分赛道与标的全景图')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

meta = doc.add_paragraph()
run = meta.add_run('基于深企投产业研究院《2025年光通信产业链研究报告》| 2026年6月4日')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
doc.add_paragraph()

# ===== 1. Industry Chain Overview =====
doc.add_heading('一、光通信全产业链架构', level=1)

doc.add_paragraph('光通信产业链分为五大环节：', style='List Bullet')
doc.add_paragraph('上游：光电芯片（光芯片 + 电芯片）', style='List Bullet')
doc.add_paragraph('中游：光器件（光无源器件 + 光有源器件）', style='List Bullet')
doc.add_paragraph('中游：光模块（发射/接收一体化封装）', style='List Bullet')
doc.add_paragraph('中游：光纤光缆', style='List Bullet')
doc.add_paragraph('下游：光通信设备', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('光通信工作原理：电信号→光发射模块(激光器)→光纤传输→光接收模块(探测器)→电信号', style='List Bullet')

doc.add_paragraph()

# ===== 2. Optical Chips =====
doc.add_heading('二、光电芯片', level=1)

doc.add_heading('2.1 光芯片', level=2)
doc.add_paragraph('光芯片是光通信系统的核心，决定传输速率和距离。主要包括激光器芯片(VCSEL/DFB/EML)和探测器芯片(PIN/APD)。')

table = doc.add_table(rows=6, cols=6)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['公司', '代码', '市值(亿)', 'PE(现)', '核心产品', 'AI受益环节']):
    table.rows[0].cells[i].text = h

chip_data = [
    ['源杰科技', '688498', '~380', '480x', 'DFB/EML激光器芯片', '100G/200G光芯片,数据中心'],
    ['长光华芯', '688048', '~80', '2037x', 'VCSEL/高功率激光芯片', '车载激光雷达+光通信'],
    ['仕佳光子', '688313', '~60', '亏损', 'PLC光分路器/AWG芯片', '数据中心AWG'],
    ['永鼎股份', '600105', '~50', 'N/A', '光芯片+光器件布局', '产业链延伸'],
    ['武汉光安', '未上市', '—', '—', '25G/50G光探测器', '国产替代核心标的'],
]
for i, row_data in enumerate(chip_data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

doc.add_heading('2.2 电芯片（Driver/TIA/CDR）', level=2)
doc.add_paragraph('电芯片负责驱动激光器、放大接收信号、时钟恢复等。国产化率最低的环节，目前主要依赖进口。')

table = doc.add_table(rows=3, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['公司', '代码', '核心产品', '竞争地位', '国产替代进展']):
    table.rows[0].cells[i].text = h

electrical = [
    ['澜起科技', '688008', 'DDR5接口/SerDes', '全球龙头', '内存接口芯片全球市占率~45%'],
    ['裕太微', '688515', '以太网PHY芯片', '国内领先', '车载以太网PHY突破'],
]
for i, row_data in enumerate(electrical):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

# ===== 3. Optical Components =====
doc.add_heading('三、光器件', level=1)

doc.add_heading('3.1 光无源器件', level=2)
doc.add_paragraph('包括光纤连接器、光分路器、波分复用器(WDM)、光隔离器、光衰减器、光开关、光滤波器等。不涉及光电转换。')

table = doc.add_table(rows=7, cols=6)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['公司', '代码', '核心产品', 'AI关联度', '竞争格局', '备注']):
    table.rows[0].cells[i].text = h

passive = [
    ['天孚通信', '300394', 'FA/MT/AWG/光引擎', '⭐⭐⭐⭐', '国内龙头', '平台化布局,800G/1.6T核心受益'],
    ['光迅科技', '002281', '全品类光器件', '⭐⭐⭐', '国内最大', '产业链最完整的光器件企业'],
    ['光库科技', '300620', '铌酸锂调制器/隔离器', '⭐⭐⭐', '细分龙头', '薄膜铌酸锂调制器前沿布局'],
    ['博创科技', '300548', 'PLC光分路器/AWG', '⭐⭐', '国内领先', 'PLC技术路线龙头'],
    ['太辰光', '300570', '光纤连接器/陶瓷插芯', '⭐⭐', '细分龙头', '全球最大的陶瓷插芯供应商'],
    ['腾景科技', '688195', '精密光学元件', '⭐⭐', '细分龙头', '高端精密光学加工'],
]
for i, row_data in enumerate(passive):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

doc.add_heading('3.2 光有源器件', level=2)
doc.add_paragraph('包括激光器(光源)、光调制器、光探测器、光放大器(EDFARaman)、光收发组件等。涉及光电转换，技术壁垒更高。')

p = doc.add_paragraph()
run = p.add_run('相关公司：')
run.bold = True
p.add_run('源杰科技(激光器芯片)、光迅科技(全品类)、华工科技(激光器+光模块)、光库科技(调制器)')
doc.add_paragraph()

# ===== 4. Optical Modules =====
doc.add_heading('四、光模块', level=1)
doc.add_paragraph('光模块是光通信产业链中AI受益最直接的环节。将光发射组件(TOSA)+光接收组件(ROSA)+电芯片封装为一体。')

doc.add_paragraph('AI光模块升级路径：400G→800G→1.6T→3.2T', style='List Bullet')
doc.add_paragraph('800G于2024年起量，1.6T预计2026年下半年开始批量交付', style='List Bullet')
doc.add_paragraph('AI数据中心对高速率光模块的需求是指数级增长', style='List Bullet')

doc.add_paragraph()

table = doc.add_table(rows=6, cols=7)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['公司', '代码', '核心产品', '客户', '800G进度', '1.6T进度', 'AI弹性']):
    table.rows[0].cells[i].text = h

module = [
    ['中际旭创', '300308', '400G/800G/1.6T', '英伟达/谷歌/Meta', '批量交付', '研发中/预计2026H2', '⭐⭐⭐⭐⭐'],
    ['新易盛', '300502', '400G/800G', '云厂商', '批量交付', '研发中', '⭐⭐⭐⭐'],
    ['联特科技', '301205', '400G/800G', '云厂商', '小批量', '早期', '⭐⭐⭐'],
    ['剑桥科技', '603083', '100G/400G/800G', '云厂商', '小批量', '早期', '⭐⭐⭐'],
    ['华工科技', '000988', '100G/400G', '数据中心+电信', '研发中', '早期', '⭐⭐'],
]
for i, row_data in enumerate(module):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

# ===== 5. Fiber Optics =====
doc.add_heading('五、光纤光缆', level=1)
doc.add_paragraph('光纤光缆是光通信的物理传输介质。AI数据中心对光纤的需求从多模转向单模，海缆/城域网也有结构性需求。')

table = doc.add_table(rows=4, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['公司', '代码', '核心产品', '产能地位', 'AI受益']):
    table.rows[0].cells[i].text = h

fiber = [
    ['长飞光纤', '601869', '预制棒/光纤/光缆', '全球第一', 'AI数据中心+海缆双驱动'],
    ['亨通光电', '600487', '光纤光缆/海缆', '国内前三', '海缆+数据中心互联'],
    ['中天科技', '600522', '光纤光缆/海缆', '国内前三', '海缆+新能源双主线'],
]
for i, row_data in enumerate(fiber):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

# ===== 6. Equipment =====
doc.add_heading('六、光通信设备', level=1)
doc.add_paragraph('光通信设备是产业链下游，包括光传输设备(OTN/WDM)、光接入设备(PON/OLT)、交换机等。')

table = doc.add_table(rows=4, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['公司', '代码', '核心产品', '竞争地位', 'AI关联']):
    table.rows[0].cells[i].text = h

equip = [
    ['中兴通讯', '000063', 'OTN/WDM/5G基站', '全球前四', 'AI算力+传输网升级'],
    ['烽火通信', '600498', 'OTN/OLT/F5G', '国内前三', '城域光传输扩容'],
    ['华为', '未上市', '全系列光设备', '全球第一', 'ICT全栈受益'],
]
for i, row_data in enumerate(equip):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

# ===== 7. Full Chain Summary =====
doc.add_heading('七、全产业链标的总表（最新行情）', level=1)

table = doc.add_table(rows=23, cols=7)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['环节', '公司', '代码', '现价', 'PE', 'FY26E净利', 'AI受益评级']):
    table.rows[0].cells[i].text = h

all_stocks = [
    ['光芯片', '源杰科技', '688498', '¥1,374', '480x', '3亿估', '⭐⭐⭐'],
    ['光芯片', '长光华芯', '688048', '¥390', '2037x', '2亿估', '⭐⭐'],
    ['光芯片', '华工科技', '000988', '¥158', '94x', '16亿', '⭐⭐⭐'],
    ['光器件', '天孚通信', '300394', '¥480', '172x', '18亿', '⭐⭐⭐⭐'],
    ['光器件', '光迅科技', '002281', '¥221', '172x', '12亿', '⭐⭐⭐'],
    ['光器件', '光库科技', '300620', '¥319', '378x', '5亿估', '⭐⭐⭐'],
    ['光器件', '博创科技', '300548', '¥230', '181x', '4亿估', '⭐⭐'],
    ['光器件', '太辰光', '300570', '¥172', '137x', '3亿估', '⭐⭐'],
    ['光器件', '腾景科技', '688195', '¥237', '595x', '1.5亿估', '⭐⭐'],
    ['光模块', '中际旭创', '300308', '¥1,280', '96x', '90亿', '⭐⭐⭐⭐⭐'],
    ['光模块', '新易盛', '300502', '¥776', '72x', '40亿', '⭐⭐⭐⭐'],
    ['光模块', '联特科技', '301205', '¥332', '493x', '6亿估', '⭐⭐⭐'],
    ['光模块', '剑桥科技', '603083', '¥204', '206x', '5亿估', '⭐⭐⭐'],
    ['光模块', '德科立', '688205', '¥234', '486x', '3.5亿估', '⭐⭐'],
    ['光纤光缆', '长飞光纤', '601869', '—', '—', '—', '⭐⭐'],
    ['光纤光缆', '亨通光电', '600487', '—', '—', '—', '⭐⭐'],
    ['光纤光缆', '中天科技', '600522', '—', '—', '—', '⭐⭐'],
    ['电芯片', '澜起科技', '688008', '—', '—', '—', '⭐⭐'],
    ['光设备', '中兴通讯', '000063', '—', '—', '—', '⭐⭐⭐'],
    ['光设备', '烽火通信', '600498', '—', '—', '—', '⭐⭐'],
    ['光设备', '华为', '未上市', '—', '—', '—', '⭐⭐⭐⭐⭐'],
    ['光设备', '英伟达', 'NVDA', '—', '—', '—', '⭐⭐⭐⭐⭐(客)'],
]
for i, row_data in enumerate(all_stocks):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

# ===== 8. Summary =====
doc.add_heading('八、投资结论', level=1)

doc.add_paragraph('按AI受益确定性排序，剔除PE泡沫因素：', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('第一梯队（最核心AI受益）：')
run.bold = True
p.add_run('中际旭创 > 新易盛 > 天孚通信')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('第二梯队（AI+国产替代）：')
run.bold = True
p.add_run('光迅科技 > 源杰科技 > 华工科技')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('第三梯队（AI间接/配套）：')
run.bold = True
p.add_run('光库科技 > 长飞光纤 > 中兴通讯 > 博创科技')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('泡沫警示（PE远超合理区间）：')
run.bold = True
p.add_run('腾景科技(595x) > 长光华芯(2037x) > 德科立(486x) > 联特科技(493x)')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('周期跟踪核心指标：')
run.bold = True
doc.add_paragraph('英伟达/谷歌/Meta资本开支指引 — 决定光模块需求的根本', style='List Bullet')
doc.add_paragraph('800G→1.6T升级节奏 — 决定产品ASP走势', style='List Bullet')
doc.add_paragraph('硅光技术进展 — 决定长期竞争格局变化', style='List Bullet')
doc.add_paragraph('光芯片国产替代率 — 决定国内产业链自主可控度', style='List Bullet')

doc.add_paragraph()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Y —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据来源：深企投产业研究院光通信研报 / 腾讯财经  分析日期：2026年6月4日')
run.font.size = Pt(9)

output = r'C:\Users\Duke Wang\.openclaw\workspace\光通信全产业链_细分赛道与标的全景_20260604.docx'
doc.save(output)
print(f'Saved: {output}')
