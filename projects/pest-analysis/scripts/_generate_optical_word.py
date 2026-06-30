#!/usr/bin/env python3
"""Generate Word: Optical communication cycle analysis."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)

title = doc.add_heading('', level=0)
run = title.add_run('光通信产业链 — AI驱动周期分析与价格预判')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
run = sub.add_run('基于光通信产业链研报 + AI发展趋势 + 实时行情')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()

# ===== Price Target Matrix =====
doc.add_heading('一、光通信14只标的估值全景', level=1)

table = doc.add_table(rows=15, cols=8)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['标的', '代码', '现价', 'PE(现)', '赛道', '合理PE', '泡沫阈值', '状态']):
    table.rows[0].cells[i].text = h

data = [
    ['中际旭创', '300308', '¥1,280', '96x', '光模块', '25x', '60x', '🔴 泡沫'],
    ['新易盛', '300502', '¥776', '72x', '光模块', '25x', '60x', '🔴 泡沫'],
    ['联特科技', '301205', '¥332', '493x', '光模块', '25x', '60x', '🔴 泡沫'],
    ['剑桥科技', '603083', '¥204', '206x', '光模块', '25x', '60x', '🔴 泡沫'],
    ['天孚通信', '300394', '¥480', '172x', '光器件', '30x', '70x', '🔴 泡沫'],
    ['光迅科技', '002281', '¥221', '172x', '光器件', '30x', '70x', '🔴 泡沫'],
    ['光库科技', '300620', '¥319', '378x', '光器件', '30x', '70x', '🔴 泡沫'],
    ['太辰光', '300570', '¥172', '137x', '光器件', '30x', '70x', '🔴 泡沫'],
    ['博创科技', '300548', '¥230', '181x', '光器件', '30x', '70x', '🔴 泡沫'],
    ['腾景科技', '688195', '¥237', '595x', '光器件', '30x', '70x', '🔴 泡沫'],
    ['源杰科技', '688498', '¥1,374', '480x', '光芯片', '50x', '150x', '🔴 泡沫'],
    ['德科立', '688205', '¥234', '486x', '光芯片', '50x', '150x', '🔴 泡沫'],
    ['长光华芯', '688048', '¥390', '2037x', '光芯片', '50x', '150x', '🔴 泡沫'],
    ['华工科技', '000988', '¥158', '94x', '光芯片', '50x', '150x', '🟡 合理区'],
]
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

# ===== Key Anchors =====
doc.add_heading('二、重点标的深度分析', level=1)

doc.add_heading('2.1 中际旭创 — AI光模块绝对龙头', level=2)
p = doc.add_paragraph()
p.add_run('核心逻辑：').bold = True
p.add_run('800G光模块放量+1.6T接力，英伟达/谷歌/Meta核心供应商。AI数据中心内部互联从400G→800G→1.6T是确定性趋势。')
p = doc.add_paragraph()
p.add_run('当前PE 95x看起来高，但FY26E净利估90亿，EPS约¥22，对应Forward PE约58x。若AI资本开支持续超预期，有望被业绩消化。').bold = False

doc.add_paragraph()

doc.add_heading('2.2 新易盛 — 弹性标的但PE偏高', level=2)
p = doc.add_paragraph()
p.add_run('核心逻辑：').bold = True
p.add_run('高速光模块，云厂商核心供应商。但PE 72x已远超合理25-45x区间。FY26E净利40亿对应Forward PE约38x。现价¥776距合理价¥431仍有44%下行空间。')

doc.add_paragraph()

doc.add_heading('2.3 天孚通信 — 光器件平台化龙头', level=2)
p = doc.add_paragraph()
p.add_run('核心逻辑：').bold = True
p.add_run('FA+AWG光器件核心供应商，受益800G/1.6T放量。但PE 172x严重泡沫化。现价¥480距合理价¥84（器件合理PE 30x）有83%下行风险。')

doc.add_paragraph()

# ===== Comparison with Storage =====
doc.add_heading('三、光通信 vs 存储芯片：泡沫对比', level=1)

table = doc.add_table(rows=3, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['对比维度', '存储芯片', '光通信', '差异', '结论']):
    table.rows[0].cells[i].text = h
table.rows[1].cells[0].text = '合理PE区间标的占比'
table.rows[1].cells[1].text = '50%（8/16只合理）'
table.rows[1].cells[2].text = '7%（1/14只合理）'
table.rows[1].cells[3].text = '光通信泡沫更大'
table.rows[1].cells[4].text = '光通信赛道更拥挤'
table.rows[2].cells[0].text = '平均PE'
table.rows[2].cells[1].text = '~80x'
table.rows[2].cells[2].text = '~380x'
table.rows[2].cells[3].text = '光通信PE高5倍'
table.rows[2].cells[4].text = '光通信需业绩消化'

doc.add_paragraph()

# ===== Investment Framework =====
doc.add_heading('四、投资框架', level=1)

doc.add_heading('4.1 赛道定位', level=2)
doc.add_paragraph('光模块(中际旭创/新易盛): AI最直接受益环节，但估值已充分定价', style='List Bullet')
doc.add_paragraph('光器件(天孚/光迅): 弹性次于模块，但目前PE同样泡沫化', style='List Bullet')
doc.add_paragraph('光芯片(源杰/德科立): 国产替代空间大，但盈利尚弱，高PE需注意', style='List Bullet')

doc.add_paragraph()

doc.add_heading('4.2 周期启动信号', level=2)
signals = [
    '英伟达/谷歌/Meta资本开支超预期增长 — 最重要先行指标',
    '800G光模块季度出货量环比增长>30% — 确认景气上行',
    '1.6T光模块进入批量交付阶段 — 下一轮增长引擎',
    '光芯片产能瓶颈缓解 — 供给端利好',
    '硅光技术突破带动成本下降 — 长期结构利好',
]
for s in signals:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph()

doc.add_heading('4.3 对比存储芯片：光通信的优劣势', level=2)
doc.add_paragraph('优势：AI确定性更高（光模块直接挂钩英伟达）、成长空间更大（1.6T才刚开始）', style='List Bullet')
doc.add_paragraph('劣势：估值泡沫更严重、业绩兑现存在预期差、竞争格局更分散', style='List Bullet')

doc.add_paragraph()

# ===== Footer =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Y —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据来源：腾讯财经 / 深企投产业研究院光通信研报  分析日期：2026年6月4日')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

output = r'C:\Users\Duke Wang\.openclaw\workspace\光通信产业链_AI周期分析与价格预判_20260604.docx'
doc.save(output)
print(f'Saved: {output}')
