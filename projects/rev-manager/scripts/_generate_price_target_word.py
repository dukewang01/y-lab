#!/usr/bin/env python3
"""Generate Word: Storage cycle price target analysis."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

title = doc.add_heading('', level=0)
run = title.add_run('存储芯片超级周期 — 价格预判与泡沫临界值分析')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
run = sub.add_run('基于FY26E业绩预测 + PE估值锚点矩阵')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ===== Price Target Matrix =====
doc.add_heading('一、价格锚点矩阵（核心关注的8只）', level=1)

table = doc.add_table(rows=9, cols=8)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['标的', '现价', 'FY26E净利', '合理价', '偏高价', '泡沫价', '现→合理', '区间状态']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

rows_data = [
    ['德明利', '¥681', '6.9亿', '¥936', '¥1,560', '¥2,184', '+37%', '🟢 合理区'],
    ['佰维存储', '¥341', '12.9亿', '¥618', '¥1,082', '¥1,391', '+81%', '🟢 合理区'],
    ['江波龙', '¥568', '13.8亿', '¥776', '¥1,330', '¥1,774', '+37%', '🟢 合理区'],
    ['北方华创', '¥627', '76.5亿', '¥749', '¥1,284', '¥1,605', '+20%', '🟢 合理区'],
    ['中微公司', '¥289', '25.0亿', '¥324', '¥607', '¥809', '+12%', '🟢 合理区'],
    ['兆易创新', '¥529', '40.5亿', '¥453', '¥755', '¥906', '-14%', '🔴 泡沫区'],
    ['长电科技', '¥80', '27.9亿', '¥29', '¥47', '¥59', '-63%', '🔴 泡沫区'],
    ['通富微电', '¥71', '14.9亿', '¥24', '¥41', '¥53', '-67%', '🔴 泡沫区'],
]
for i, row_data in enumerate(rows_data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

# ===== Category Breakdown =====
doc.add_heading('二、分赛道详细估值', level=1)

# Module section
doc.add_heading('2.1 模组三杰（最安全标的）', level=2)
p = doc.add_paragraph('德明利(001309)、佰维存储(688525)、江波龙(301308)')
p = doc.add_paragraph('当前PE 38-44x，处于历史合理区间偏下。在DRAM/NAND涨价的超级周期中，模组企业直接受益于库存增值和价差扩大。2026年业绩增长确定性最高。')
doc.add_paragraph('合理价：对应PE 30-40x，较现价仍有37-81%上行空间', style='List Bullet')
doc.add_paragraph('偏高警戒价：对应PE 50-70x，到达则关注基本面是否配合', style='List Bullet')
doc.add_paragraph('泡沫出货价：对应PE 70-90x，到达即减仓', style='List Bullet')

doc.add_paragraph()

# Equipment
doc.add_heading('2.2 设备双雄（中线国产替代逻辑）', level=2)
p = doc.add_paragraph('北方华创(002371)、中微公司(688012)')
p = doc.add_paragraph('PE 82-99x，虽不低但有国产替代长逻辑支撑。存储厂扩产周期直接拉动设备订单。')
doc.add_paragraph('合理价¥749/¥324，现价仍有12-20%上行空间', style='List Bullet')
doc.add_paragraph('国产替代赋予更高PE容忍度，泡沫阈值150-200x', style='List Bullet')

doc.add_paragraph()

# Bubble
doc.add_heading('2.3 泡沫区标的（需警惕）', level=2)
p = doc.add_paragraph('兆易创新(603986): PE 129x已超历史峰值100x、泡沫阈值120x。FY26E净利40.5亿可消化部分估值，但现价¥529距合理价¥453仍有14%下行空间。')
p = doc.add_paragraph('长电科技(600584): PE 87x，封测行业合理PE仅25-40x。现价¥80远超泡沫价¥59，任何周期转弱信号都会导致剧烈回调。')
p = doc.add_paragraph('通富微电(002156): PE 75x > 泡沫阈值45x，同上逻辑。')

doc.add_paragraph()

# ===== Investment Framework =====
doc.add_heading('三、投资框架', level=1)

doc.add_heading('3.1 仓位配置建议', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['类型', '配置比例', '首选标的']):
    table.rows[0].cells[i].text = h
table.rows[1].cells[0].text = '保守型'
table.rows[1].cells[1].text = '70%+合理区标的'
table.rows[1].cells[2].text = '德明利/佰维/江波龙/北方华创'
table.rows[2].cells[0].text = '均衡型'
table.rows[2].cells[1].text = '50%合理+30%偏高'
table.rows[2].cells[2].text = '上述+中微/雅克/安集'
table.rows[3].cells[0].text = '激进型'
table.rows[3].cells[1].text = '<40%合理区,配合止损设置'
table.rows[3].cells[2].text = '不可重仓泡沫区标的'

doc.add_paragraph()

doc.add_heading('3.2 周期反转信号清单', level=2)
signals = [
    'DRAM/NAND现货价格连续2周下跌 — 确认拐点',
    'SK海力士/三星资本开支计划缩减 — 供给端信号',
    'CSP云厂商资本开支增速放缓 — 需求端信号',
    '存储模组厂毛利率见顶回落 — 行业利润转移',
    '下游客户库存周转天数上升 — 供需失衡前兆',
]
for s in signals:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph()

doc.add_heading('3.3 关键触发条件', level=2)
doc.add_paragraph('触发减仓：出现上述信号3条以上 + 持仓标的到达泡沫价', style='List Bullet')
doc.add_paragraph('触发清仓：DRAM/NAND现货价连续1个月下跌 + 行业龙头预告业绩不及预期', style='List Bullet')
doc.add_paragraph('触发加仓：周期信号未触发 + 优质标的PE回落至合理价以下', style='List Bullet')

doc.add_paragraph()

# ===== Footer =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Y —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据来源：腾讯财经 / 起点财经 / 浙商证券研报  分析日期：2026年6月4日')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output = r'C:\Users\Duke Wang\.openclaw\workspace\存储芯片周期_价格预判与泡沫分析_20260604.docx'
doc.save(output)
print(f'Saved: {output}')
