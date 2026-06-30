#!/usr/bin/env python3
"""Generate combined Word report from two storage chip reports."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

# ===== Title =====
title = doc.add_heading('', level=0)
run = title.add_run('2026年全球存储芯片行业深度分析报告')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
run = sub.add_run('融合两份研报：行业趋势 + 市场格局 + 国产厂商布局')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

meta = doc.add_paragraph()
run = meta.add_run('数据来源：起点财经 / 行业研究平台  |  2026年6月 | 共95页合并研究')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()

# ===== 1. Market Overview =====
doc.add_heading('一、市场总览：超级周期到来', level=1)

doc.add_paragraph('全球存储芯片市场正经历前所未有的"超级周期"。2025年市场规模突破2000亿美元，预计2026年攀升至约5516亿美元，同比增长134%。')

table = doc.add_table(rows=3, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['指标', '2025年', '2026年(预计)', '变化']):
    table.rows[0].cells[i].text = h
table.rows[1].cells[0].text = '全球市场规模'
table.rows[1].cells[1].text = '>2000亿美元'
table.rows[1].cells[2].text = '~5516亿美元'
table.rows[1].cells[3].text = '+134%'
table.rows[2].cells[0].text = '核心驱动'
table.rows[2].cells[1].text = 'AI需求爆发'
table.rows[2].cells[2].text = 'AI+消费升级+国产替代'
table.rows[2].cells[3].text = '结构性增长'

doc.add_paragraph()
doc.add_heading('1.1 三大驱动力', level=2)
doc.add_paragraph('AI算力需求爆发：生成式AI对存储的需求是传统服务器的8-10倍，NAND需求提升12倍以上。', style='List Bullet')
doc.add_paragraph('消费电子升级换代：智能手机、PC、AIPC对存储容量和性能要求持续提升。', style='List Bullet')
doc.add_paragraph('国产替代加速：地缘政治背景下，本土企业加速突破技术瓶颈，国产化率持续提升。', style='List Bullet')

# ===== 2. Market Structure =====
doc.add_heading('二、市场结构：DRAM与NAND Flash主导', level=1)

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['品类', '市场份额', '核心应用']):
    table.rows[0].cells[i].text = h
table.rows[1].cells[0].text = 'DRAM'
table.rows[1].cells[1].text = '56%'
table.rows[1].cells[2].text = '计算机内存、服务器、移动设备、AI服务器核心部件'
table.rows[2].cells[0].text = 'NAND Flash'
table.rows[2].cells[1].text = '38-40%'
table.rows[2].cells[2].text = 'SSD、智能手机、平板电脑、嵌入式系统'
table.rows[3].cells[0].text = '其他(NOR/EEPROM等)'
table.rows[3].cells[1].text = '~4%'
table.rows[3].cells[2].text = '物联网、汽车电子、工业控制'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('HBM（高带宽内存）：')
run.bold = True
p.add_run('AI时代的核心增长极。2025年市场规模超120亿美元，从HBM3(1.6TB/s/32GB)向HBM4(2TB/s/64GB)演进。SK海力士投资130亿美元建设全球最大HBM工厂。')

# ===== 3. AI Demand =====
doc.add_heading('三、AI驱动的需求变革', level=1)

doc.add_paragraph('AI算力需求正从多个维度重塑存储芯片市场：')

items = [
    ('服务器端', 'CSPs资本开支持续上修。2026年全球八大CSPs资本开支预计超2500亿美元，AI服务器出货量渗透率预计2029年超30%。英伟达BlueField-4新型AI原生存储基础设施加速代理式AI。'),
    ('手机端', 'AI大模型本地化部署驱动LPDDR+UFS需求增长。AIGC手机渗透率2029年预计超80%。DRAM单机容量从8GB向12-16GB升级。'),
    ('PC端', 'AIPC加速存储容量扩张。AIPC渗透率2029年预计超60%。PC存储配置从16GB+512GB向32GB+1TB升级。'),
    ('企业级', 'AI服务器DRAM容量和成本大幅提升。HBM从高端利基产品演变为AI芯片生态"命脉"。企业级SSD需求激增。'),
]
for title, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(title + '：')
    run.bold = True
    p.add_run(desc)

# ===== 4. Competitive Landscape =====
doc.add_heading('四、竞争格局：三巨头垄断与国产崛起', level=1)

doc.add_heading('4.1 全球竞争格局', level=2)

table = doc.add_table(rows=4, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['厂商', 'DRAM份额', 'NAND份额', '战略方向', '核心优势']):
    table.rows[0].cells[i].text = h
table.rows[1].cells[0].text = '三星'
table.rows[1].cells[1].text = '~30%'
table.rows[1].cells[2].text = '35%'
table.rows[1].cells[3].text = '全品类布局'
table.rows[1].cells[4].text = '3D V-NAND技术领先'
table.rows[2].cells[0].text = 'SK海力士'
table.rows[2].cells[1].text = '36.7%(第一)'
table.rows[2].cells[2].text = '~25%'
table.rows[2].cells[3].text = '聚焦HBM'
table.rows[2].cells[4].text = 'HBM领域技术领先，DRAM第一'
table.rows[3].cells[0].text = '美光'
table.rows[3].cells[1].text = '~25%'
table.rows[3].cells[2].text = '~20%'
table.rows[3].cells[3].text = '退出消费级，专注企业级'
table.rows[3].cells[4].text = 'HBM3E及HBM4量产'

doc.add_paragraph()
doc.add_heading('4.2 国产力量崛起', level=2)

table = doc.add_table(rows=3, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['国产厂商', '品类', '市场份额', '技术突破', '客户绑定']):
    table.rows[0].cells[i].text = h
table.rows[1].cells[0].text = '长江存储(YMTC)'
table.rows[1].cells[1].text = 'NAND Flash'
table.rows[1].cells[2].text = '14%(2025年)'
table.rows[1].cells[3].text = 'Xtacking架构，232层3D NAND'
table.rows[1].cells[4].text = '阿里云、腾讯云、华为、小米'
table.rows[2].cells[0].text = '长鑫存储(CXMT)'
table.rows[2].cells[1].text = 'DRAM'
table.rows[2].cells[2].text = '8-12%(2025年)'
table.rows[2].cells[3].text = '19nm DDR4良率>90%，HBM3样品交付'
table.rows[2].cells[4].text = '华为、小米、国内手机/PC厂商'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('关键数据：')
run.bold = True
p.add_run('长江存储NAND份额从年初8%提升至14%，年产量达273万片。长鑫存储合肥基地月产能从12万片扩至18万片，单位晶圆成本较韩国厂商低15-20%。')

# ===== 5. Chinese Storage Companies =====
doc.add_heading('五、国产存储产业链重点公司分析', level=1)

companies = [
    ('德明利', '主控芯片自研为核心的模组专家', 
     '存储模组企业，主控芯片自研为核心竞争力。聚焦高附加值场景。切入企业级存储，构建智能化工厂。'),
    ('佰维存储', '"研发封测一体化"优势',
     '六大产品条线覆盖全场景。授权品牌覆盖度超越国内同行。AI眼镜存储、嵌入式存储等多领域布局。深度绑定终端抢占AI机遇。'),
    ('江波龙', '全产业链布局的存储平台领军',
     '企业级突破与高端化驱动成长。自研主控芯片UFS 4.1。旗下品牌矩阵协同效应。四大产品线完备，研发投入持续加大。'),
    ('神工股份', '电子级单晶硅材料龙头',
     '刻蚀用单晶硅材料核心供应商。硅零部件产品丰富。深度受益存储大周期。'),
    ('雅克科技', '国内领先前驱体厂商',
     '前驱体产品进入海力士、长鑫供应链。中国半导体前驱体市场持续扩张。AI存储需求带动材料环节增长。'),
    ('兆易创新', '平台化布局的存储设计龙头',
     'A股唯一覆盖三大存储品类（NOR/NAND/DRAM）。全球NOR Flash市占率18.5%（国内第一）。车规级存储占比达35%，切入特斯拉、比亚迪供应链。2025年Q3净利润同比增长68%。'),
    ('香农芯创', '"分销+产品"双轮驱动',
     '与SK海力士设立海普存储聚焦企业级存储。稀缺原厂资源赋能国产化突围。生态合作伙伴广泛。'),
    ('普冉股份', 'NOR Flash新星',
     '"存储+"战略打开新成长空间。产品覆盖消费/工业/汽车多领域。营收与净利润持续增长。'),
    ('聚辰股份', '全球第三大SPD供应商',
     'EEPROM芯片汽车电子应用广泛。音圈马达驱动芯片布局。存储与驱动双赛道扩张。'),
    ('东芯股份', '利基存储+存算联布局',
     '利基存储基本盘稳固。"存算联"一体化布局打开长期空间。投资上海砺算布局GPU。架构创新解决冯诺依曼瓶颈。全球Wi-Fi芯片组市场增长受益。'),
]

for name, tagline, desc in companies:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}')
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(f'定位：{tagline}', style='List Bullet')
    doc.add_paragraph(f'要点：{desc}', style='List Bullet')

# ===== 6. Supply Chain =====
doc.add_heading('六、上游设备材料国产化', level=1)

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['环节', '代表企业', '核心产品', '国产化进展']):
    table.rows[0].cells[i].text = h
table.rows[1].cells[0].text = '刻蚀设备'
table.rows[1].cells[1].text = '中微公司'
table.rows[1].cells[2].text = 'ICP刻蚀设备'
table.rows[1].cells[3].text = '覆盖DRAM/3D NAND 95%刻蚀场景'
table.rows[2].cells[0].text = '薄膜沉积'
table.rows[2].cells[1].text = '北方华创'
table.rows[2].cells[2].text = 'PVD/CVD设备'
table.rows[2].cells[3].text = '12英寸外延设备全覆盖'
table.rows[3].cells[0].text = '抛光材料'
table.rows[3].cells[1].text = '安集科技'
table.rows[3].cells[2].text = '存储芯片抛光液'
table.rows[3].cells[3].text = '3D NAND/DRAM适配，核心供应商'
table.rows[4].cells[0].text = '大尺寸硅片'
table.rows[4].cells[1].text = '沪硅产业'
table.rows[4].cells[2].text = '300mm硅片'
table.rows[4].cells[3].text = '已量产，供长江存储/长鑫存储'

doc.add_paragraph()
doc.add_paragraph('EDA/IP工具：华大九天（存储电路全定制设计）、概伦电子（仿真良率控制）、广立微（良率优化）。', style='List Bullet')
doc.add_paragraph('国家大基金三期2500亿元精准扶持，加速设备材料国产化。', style='List Bullet')

# ===== 7. Technology Trends =====
doc.add_heading('七、技术发展趋势', level=1)

doc.add_heading('7.1 HBM技术演进：从HBM3到HBM4', level=2)
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['指标', 'HBM3', 'HBM4']):
    table.rows[0].cells[i].text = h
table.rows[1].cells[0].text = '带宽'
table.rows[1].cells[1].text = '最高1.6 TB/s'
table.rows[1].cells[2].text = '2 TB/s'
table.rows[2].cells[0].text = '单堆栈容量'
table.rows[2].cells[1].text = '最高32GB'
table.rows[2].cells[2].text = '最高64GB'

doc.add_paragraph()
doc.add_heading('7.2 存算一体技术(CIM)', level=2)
doc.add_paragraph('将计算单元直接嵌入存储单元，实现"数据不移动，计算在存储中进行"，能效比提升数倍，从根本上解决冯·诺依曼瓶颈。从实验室加速走向产业化，成为下一代AI芯片和边缘计算设备的核心解决方案。')

doc.add_heading('7.3 先进封装革命', level=2)
doc.add_paragraph('Chiplet、2.5D/3D堆叠技术成为提升芯片性能和集成度的关键。封装环节从"后道工序"跃升为"核心战场"。SK海力士130亿美元HBM工厂验证了这一趋势。')

# ===== 8. Investment =====
doc.add_heading('八、投资机会与赛道', level=1)

doc.add_paragraph('四大投资赛道：')
invest_opts = [
    ('技术领先者', '长江存储(NAND)、兆易创新(NOR/DRAM/车规)、长鑫存储(DRAM/HBM)'),
    ('设备材料国产化', '中微公司(刻蚀)、北方华创(薄膜沉积)、沪硅产业(硅片)'),
    ('先进封测', '长电科技、通富微电、深科技（受益Chiplet/3D堆叠/HBM封装）'),
    ('生态协同', '华大九天(EDA)、江波龙(模组)、佰维存储(封测一体化)'),
]
for title, desc in invest_opts:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}：')
    run.bold = True
    p.add_run(desc)

# ===== 9. Future =====
doc.add_heading('九、未来展望（2026-2030）', level=1)

table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['年份', '关键里程碑']):
    table.rows[0].cells[i].text = h
table.rows[1].cells[0].text = '2026'
table.rows[1].cells[1].text = 'AI算力需求持续爆发式增长，存储市场达5516亿美元'
table.rows[2].cells[0].text = '2027'
table.rows[2].cells[1].text = '存算一体技术产业化，国产自给率目标40%'
table.rows[3].cells[0].text = '2028'
table.rows[3].cells[1].text = 'HBM4大规模应用，国产厂商跻身全球前列'
table.rows[4].cells[0].text = '2030'
table.rows[4].cells[1].text = '存算一体成为主流，全产业链生态竞争定格局'

doc.add_paragraph()
doc.add_paragraph('核心结论：', style='List Bullet')
doc.add_paragraph('AI驱动的结构性高景气将贯穿2026-2030年，市场规模预计突破5500亿美元。', style='List Bullet')
doc.add_paragraph('技术革命（存算一体、HBM4、先进封装）重塑产业逻辑，中国存储从"跟跑"迈向"并跑"乃至"领跑"。', style='List Bullet')
doc.add_paragraph('国产化改写全球格局，长江存储/长鑫存储为代表的本土企业正从"配角"迈向"主角"。', style='List Bullet')
doc.add_paragraph('生态化竞争成为关键胜负手——谁能构建完整的"技术+产能+生态"体系，谁将主导未来。', style='List Bullet')

doc.add_paragraph()

# ===== Footer =====
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— Y —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据来源：起点财经 / 投研报告平台\n整理日期：2026年6月4日')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output = r'C:\Users\Duke Wang\.openclaw\workspace\2026年全球存储芯片行业深度分析报告_合并版.docx'
doc.save(output)
print(f'Saved: {output}')
print(f'Size: {os.path.getsize(output)} bytes')
